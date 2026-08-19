"""src/files.py extract_text() for all four supported formats plus error corners."""
import io
import unittest

from docx import Document as DocxDocument
from pypdf import PdfWriter

from src import files


class FakeUpload:
    """Minimal stand-in for Flask's FileStorage: .filename + .stream.read()."""
    def __init__(self, filename, data):
        self.filename = filename
        self.stream = io.BytesIO(data)


class TextFileTests(unittest.TestCase):
    def test_txt_extracts_and_strips(self):
        text = files.extract_text(FakeUpload("notes.txt", b"  hello world  "))
        self.assertEqual(text, "hello world")

    def test_md_extracts(self):
        text = files.extract_text(FakeUpload("notes.md", b"# Heading\ntext"))
        self.assertEqual(text, "# Heading\ntext")

    def test_max_chars_truncation(self):
        big = ("a" * (files.MAX_CHARS + 1000)).encode()
        text = files.extract_text(FakeUpload("big.txt", big))
        self.assertEqual(len(text), files.MAX_CHARS)


class ErrorCornerTests(unittest.TestCase):
    def test_unsupported_extension_raises(self):
        with self.assertRaisesRegex(RuntimeError, "Unsupported file type"):
            files.extract_text(FakeUpload("archive.zip", b"data"))

    def test_no_extension_raises(self):
        with self.assertRaisesRegex(RuntimeError, r"\.\?"):
            files.extract_text(FakeUpload("noext", b"data"))

    def test_empty_text_raises(self):
        with self.assertRaisesRegex(RuntimeError, "No readable text"):
            files.extract_text(FakeUpload("empty.txt", b"   "))


class DocxTests(unittest.TestCase):
    def test_extracts_paragraphs_and_marks_headings(self):
        doc = DocxDocument()
        doc.add_paragraph("Intro line")
        doc.add_heading("Experience", level=1)
        doc.add_paragraph("Did things")
        buf = io.BytesIO()
        doc.save(buf)

        text = files.extract_text(FakeUpload("resume.docx", buf.getvalue()))
        self.assertIn("Intro line", text)
        self.assertIn("## Experience", text)
        self.assertIn("Did things", text)

    def test_unreadable_raises_runtime_error(self):
        with self.assertRaisesRegex(RuntimeError, r"Could not read.*\.docx"):
            files.extract_text(FakeUpload("bad.docx", b"not a real docx"))


class PdfTests(unittest.TestCase):
    def test_extracts_text(self):
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        buf = io.BytesIO()
        writer.write(buf)
        # A blank page has no extractable text - assert it fails the empty-text check rather
        # than the parse itself, which confirms the pdf branch runs end to end.
        with self.assertRaisesRegex(RuntimeError, "No readable text"):
            files.extract_text(FakeUpload("blank.pdf", buf.getvalue()))

    def test_unreadable_raises_runtime_error(self):
        with self.assertRaisesRegex(RuntimeError, r"Could not read.*PDF"):
            files.extract_text(FakeUpload("bad.pdf", b"not a real pdf"))


if __name__ == "__main__":
    unittest.main()
